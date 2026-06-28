"""Export proposals to Word (.docx) and PDF (WeasyPrint with fpdf2 fallback)."""

from __future__ import annotations

import io
import re
from datetime import date
from html.parser import HTMLParser
from typing import Any

from docx.shared import Inches, Pt
from proposal_defaults import PROPOSAL_SECTIONS, SECTION_TITLES

COMPANY_SHORT = "Northern RE Investments LLC"
COMPANY_LEGAL = "Northern RE Investments, LLC"
NAVY = "#1e3a5f"


def _today_str() -> str:
    return date.today().strftime("%B %d, %Y")


def _date_file() -> str:
    return date.today().strftime("%Y%m%d")


def _css_escape(value: str) -> str:
    return (value or "").replace("\\", "\\\\").replace("'", "\\'")


def _sanitize_filename(part: str) -> str:
    cleaned = re.sub(r"[^\w\-.]+", "_", (part or "").strip())
    return cleaned.strip("_") or "Proposal"


def proposal_filenames(solicitation_number: str) -> dict[str, str]:
    sol = _sanitize_filename(solicitation_number)
    d = _date_file()
    return {
        "docx": f"{sol}_NorthernREInvestments_Proposal_{d}.docx",
        "pdf": f"{sol}_NorthernREInvestments_Proposal_{d}.pdf",
        "capability": f"NorthernREInvestments_CapabilityStatement_{d}.pdf",
    }


def _solicitation_number(proposal: Any) -> str:
    config = proposal.config_json if isinstance(proposal.config_json, dict) else {}
    sec_a = config.get("section_a") or {}
    return (
        str(sec_a.get("solicitation_number") or "")
        or (proposal.contract.notice_id if proposal.contract else "")
        or "Proposal"
    )


def resolve_sections(proposal: Any, sections_override: dict[str, str] | None) -> dict[str, str]:
    from proposal_service import parse_sections_from_html

    if sections_override:
        return {k: v for k, v in sections_override.items() if v}
    if proposal.sections_json:
        return dict(proposal.sections_json)
    if proposal.proposal_html:
        return parse_sections_from_html(proposal.proposal_html)
    return {}


def export_meta(proposal: Any) -> dict[str, Any]:
    config = proposal.config_json if isinstance(proposal.config_json, dict) else {}
    sec_a = config.get("section_a") or {}
    sec_b = config.get("section_b") or {}
    sol = _solicitation_number(proposal)
    return {
        "solicitation_number": sol,
        "filenames": proposal_filenames(sol),
        "date_display": _today_str(),
        "owner": sec_b,
        "section_a": sec_a,
    }


class _HtmlToDocxParser(HTMLParser):
    """Minimal HTML → python-docx paragraph/run converter."""

    def __init__(self, document, *, in_table: bool = False):
        super().__init__()
        self.document = document
        self.in_table = in_table
        self._table: list[list[list[tuple[str, bool]]]] | None = None
        self._row: list[list[tuple[str, bool]]] | None = None
        self._cell: list[tuple[str, bool]] | None = None
        self._para = None
        self._bold = False
        self._italic = False
        self._list_depth = 0
        self._pending_tables: list[list[list[str]]] = []

    def _ensure_para(self):
        if self._para is None:
            self._para = self.document.add_paragraph()
        return self._para

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "br":
            self._ensure_para().add_run("\n")
        elif tag in ("p", "div", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._para = None
            if tag.startswith("h"):
                self._para = self.document.add_paragraph()
                for run in self._para.runs:
                    pass
        elif tag == "li":
            self._para = self.document.add_paragraph(style="List Bullet")
        elif tag == "table":
            self._table = []
            self._row = None
        elif tag == "tr":
            self._row = []
        elif tag in ("td", "th"):
            self._cell = []

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag in ("p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li"):
            self._para = None
        elif tag in ("td", "th"):
            if self._row is not None and self._cell is not None:
                text = "".join(t for t, _ in self._cell).strip()
                self._row.append(text)
            self._cell = None
        elif tag == "tr":
            if self._table is not None and self._row is not None:
                self._table.append(self._row)
            self._row = None
        elif tag == "table":
            if self._table:
                self._pending_tables.append(self._table)
            self._table = None

    def handle_data(self, data):
        if not data:
            return
        if self._cell is not None:
            self._cell.append((data, self._bold))
            return
        para = self._ensure_para()
        run = para.add_run(data)
        run.bold = self._bold
        run.italic = self._italic

    def flush_tables(self):
        from docx.shared import Pt

        for table_data in self._pending_tables:
            if not table_data:
                continue
            cols = max(len(r) for r in table_data)
            table = self.document.add_table(rows=len(table_data), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(table_data):
                for ci in range(cols):
                    text = row[ci] if ci < len(row) else ""
                    cell = table.rows[ri].cells[ci]
                    cell.text = text
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)


def _add_page_number_field(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _clear_header_footer_block(block):
    for p in list(block.paragraphs):
        p._element.getparent().remove(p._element)


def _setup_docx_headers_footers(document, solicitation_number: str, date_display: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        _clear_header_footer_block(section.header)
        hp = section.header.add_paragraph()
        left = hp.add_run(COMPANY_SHORT)
        left.font.name = "Times New Roman"
        left.font.size = Pt(10)
        hp.add_run("\t" * 6)
        right = hp.add_run(solicitation_number)
        right.font.name = "Times New Roman"
        right.font.size = Pt(10)

        _clear_header_footer_block(section.footer)
        fp = section.footer.add_paragraph()
        lrun = fp.add_run("Proprietary and Confidential")
        lrun.font.name = "Times New Roman"
        lrun.font.size = Pt(9)
        fp.add_run("\t")
        _add_page_number_field(fp)
        fp.add_run("\t")
        drun = fp.add_run(date_display)
        drun.font.name = "Times New Roman"
        drun.font.size = Pt(9)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _strip_outer_html(html: str) -> str:
    text = (html or "").strip()
    text = re.sub(r"^<h[12][^>]*>.*?</h[12]>", "", text, count=1, flags=re.I | re.S)
    return text.strip()


def _add_html_to_docx(document, html: str, *, heading: str | None = None, is_cover: bool = False):
    if heading:
        h = document.add_paragraph()
        hr = h.add_run(heading)
        hr.bold = True
        hr.font.name = "Times New Roman"
        hr.font.size = Pt(14)

    parser = _HtmlToDocxParser(document)
    parser.feed(_strip_outer_html(html))
    parser.flush_tables()

    if is_cover:
        document.add_page_break()


def build_proposal_docx(proposal: Any, sections: dict[str, str], meta: dict[str, Any]) -> bytes:
    from docx import Document

    document = Document()
    _setup_docx_headers_footers(document, meta["solicitation_number"], meta["date_display"])

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    first = True
    for key in PROPOSAL_SECTIONS:
        html = sections.get(key)
        if not html:
            continue
        if not first:
            document.add_page_break()
        first = False
        title = SECTION_TITLES.get(key, key)
        is_cover = key == "cover_letter"
        _add_html_to_docx(document, html, heading=title, is_cover=is_cover)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _proposal_print_html(sections: dict[str, str], meta: dict[str, Any]) -> str:
    sol = _css_escape(meta["solicitation_number"])
    date_display = _css_escape(meta["date_display"])
    parts = [
        "<!DOCTYPE html><html><head><meta charset='utf-8'>",
        "<style>",
        "@page { size: letter; margin: 1in 1in 1.1in 1in; ",
        f"  @top-left {{ content: '{COMPANY_SHORT}'; font-family: 'Times New Roman', Times, serif; font-size: 10pt; }}",
        f"  @top-right {{ content: '{sol}'; font-family: 'Times New Roman', Times, serif; font-size: 10pt; }}",
        "  @bottom-left { content: 'Proprietary and Confidential'; font-family: 'Times New Roman', Times, serif; font-size: 9pt; }",
        "  @bottom-center { content: counter(page); font-family: 'Times New Roman', Times, serif; font-size: 9pt; }",
        f"  @bottom-right {{ content: '{date_display}'; font-family: 'Times New Roman', Times, serif; font-size: 9pt; }}",
        "}",
        "body { font-family: 'Times New Roman', Times, serif; font-size: 12pt; line-height: 1.35; color: #111; }",
        "h2.section-title { font-size: 14pt; font-weight: bold; margin: 0 0 12pt 0; bookmark-level: 1; }",
        ".section-block { page-break-before: always; }",
        ".section-block:first-child { page-break-before: auto; }",
        ".cover-letter { page-break-after: always; min-height: 9in; }",
        "table { border-collapse: collapse; width: 100%; margin: 12pt 0; }",
        "th, td { border: 1px solid #000; padding: 6px 10px; text-align: left; font-size: 12pt; }",
        "th { font-weight: bold; background: #f5f5f5; }",
        "p { margin: 0 0 10pt 0; }",
        "ul, ol { margin: 0 0 10pt 20pt; }",
        "</style></head><body>",
    ]

    first = True
    for key in PROPOSAL_SECTIONS:
        html = sections.get(key)
        if not html:
            continue
        title = SECTION_TITLES.get(key, key)
        cls = "section-block"
        if key == "cover_letter":
            cls += " cover-letter"
        break_attr = "" if first else ' style="page-break-before: always;"'
        first = False
        anchor = key.replace("_", "-")
        parts.append(f'<div class="{cls}" id="{anchor}"{break_attr}>')
        parts.append(f'<h2 class="section-title" id="bookmark-{anchor}">{title}</h2>')
        parts.append(_strip_outer_html(html))
        parts.append("</div>")

    parts.append("</body></html>")
    return "".join(parts)


def _html_to_plain(html: str) -> str:
    text = re.sub(r"<br\s*/?>", "\n", html or "", flags=re.I)
    text = re.sub(r"</p>", "\n\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _weasyprint_available() -> bool:
    try:
        import weasyprint  # noqa: F401

        return True
    except Exception:
        return False


def build_proposal_pdf_weasyprint(html: str) -> bytes:
    from weasyprint import HTML

    return HTML(string=html).write_pdf(optimize_size=("fonts", "images"))


def _fpdf_bytes(pdf) -> bytes:
    out = pdf.output()
    if isinstance(out, bytes):
        return out
    if isinstance(out, bytearray):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def build_proposal_pdf_fpdf(proposal: Any, sections: dict[str, str], meta: dict[str, Any]) -> bytes:
    from fpdf import FPDF

    pdf = FPDF(format="Letter")
    pdf.set_auto_page_break(auto=True, margin=72)
    pdf.set_margins(72, 72, 72)
    sol = meta["solicitation_number"]
    date_display = meta["date_display"]

    for key in PROPOSAL_SECTIONS:
        html = sections.get(key)
        if not html:
            continue
        title = SECTION_TITLES.get(key, key)
        pdf.add_page()
        pdf.set_font("Times", "B", 10)
        pdf.cell(0, 8, COMPANY_SHORT, align="L")
        pdf.cell(0, 8, sol, align="R", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
        pdf.set_font("Times", "B", 14)
        pdf.multi_cell(0, 8, title)
        pdf.ln(2)
        pdf.set_font("Times", "", 12)
        body = _html_to_plain(_strip_outer_html(html))
        pdf.multi_cell(0, 6, body)
        pdf.set_y(-36)
        pdf.set_font("Times", "", 9)
        pdf.cell(60, 8, "Proprietary and Confidential", align="L")
        pdf.cell(70, 8, f"Page {pdf.page_no()}", align="C")
        pdf.cell(60, 8, date_display, align="R")

    return _fpdf_bytes(pdf)


def build_proposal_pdf(proposal: Any, sections: dict[str, str], meta: dict[str, Any]) -> tuple[bytes, str]:
    """Returns (pdf_bytes, engine_used)."""
    html = _proposal_print_html(sections, meta)
    if _weasyprint_available():
        try:
            return build_proposal_pdf_weasyprint(html), "weasyprint"
        except Exception:
            pass
    return build_proposal_pdf_fpdf(proposal, sections, meta), "fpdf2"


def _capability_html(proposal: Any, sections: dict[str, str], meta: dict[str, Any]) -> str:
    owner = meta.get("owner") or {}
    cap_html = sections.get("capability_statement") or ""
    past_html = sections.get("past_performance") or ""
    cap_body = _html_to_plain(_strip_outer_html(cap_html))
    past_body = _html_to_plain(_strip_outer_html(past_html)) or (owner.get("commercial_experience") or "")

    addr_parts = [
        owner.get("address_line_1"),
        owner.get("address_line_2"),
        " ".join(
            p
            for p in [owner.get("city"), owner.get("state"), owner.get("zip")]
            if p
        ),
    ]
    address = ", ".join(p for p in addr_parts if p)

    company_data = [
        ("CAGE Code", owner.get("cage_code") or "—"),
        ("UEI", owner.get("uei") or "—"),
        ("EIN", owner.get("ein") or "—"),
        ("NAICS", f"{owner.get('primary_naics_code', '561720')} — Janitorial Services"),
        ("Business Type", "Small Business"),
        ("Structure", "Limited Liability Company"),
        ("State", owner.get("business_state") or "Wyoming"),
        ("SAM Registration", f"Active through {owner.get('sam_expiration') or '—'}"),
    ]

    cert = owner.get("certifications") or "SAM.gov registered small business."

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>
@page {{ size: letter; margin: 0.5in; }}
body {{ font-family: 'Times New Roman', Times, serif; font-size: 11pt; margin: 0; color: #111; }}
.banner-title {{ font-size: 26pt; font-weight: bold; margin: 0; padding: 16pt 0 8pt 0; }}
.navy-bar {{ background: {NAVY}; height: 8pt; margin-bottom: 18pt; }}
.columns {{ display: table; width: 100%; table-layout: fixed; }}
.col {{ display: table-cell; width: 50%; vertical-align: top; padding-right: 16pt; }}
.col:last-child {{ padding-right: 0; padding-left: 16pt; }}
h3 {{ font-size: 12pt; color: {NAVY}; margin: 0 0 8pt 0; text-transform: uppercase; letter-spacing: 0.5pt; }}
p {{ margin: 0 0 8pt 0; line-height: 1.35; }}
.data-row {{ margin-bottom: 4pt; }}
.data-label {{ font-weight: bold; }}
</style></head><body>
<div class="banner-title">{COMPANY_LEGAL}</div>
<div class="navy-bar"></div>
<div class="columns">
  <div class="col">
    <h3>Core Competencies</h3>
    <p>{cap_body[:1200] if cap_body else "Facilities support services and custodial contract management."}</p>
    <h3>Past Performance</h3>
    <p>{past_body[:900] if past_body else "Commercial field operations and retail management experience."}</p>
  </div>
  <div class="col">
    <h3>Company Data</h3>
    {''.join(f'<div class="data-row"><span class="data-label">{k}:</span> {v}</div>' for k, v in company_data)}
    <h3>Certifications</h3>
    <p>{cert}</p>
    <h3>Contact Info</h3>
    <p>{owner.get('owner_name', 'Mark Graham II')}, {owner.get('owner_title', 'Owner')}<br>
    {address}<br>
    {owner.get('business_phone', '')}<br>
    {owner.get('business_email', '')}</p>
  </div>
</div>
</body></html>"""


def build_capability_pdf_fpdf(proposal: Any, sections: dict[str, str], meta: dict[str, Any]) -> bytes:
    from fpdf import FPDF

    owner = meta.get("owner") or {}
    cap_body = _html_to_plain(_strip_outer_html(sections.get("capability_statement") or ""))
    past_body = _html_to_plain(_strip_outer_html(sections.get("past_performance") or ""))

    pdf = FPDF(format="Letter")
    pdf.set_margins(36, 36, 36)
    pdf.add_page()
    pdf.set_font("Times", "B", 24)
    pdf.cell(0, 14, COMPANY_LEGAL, new_x="LMARGIN", new_y="NEXT")
    pdf.set_fill_color(30, 58, 95)
    pdf.cell(0, 4, "", fill=True, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / 2 - 4
    y0 = pdf.get_y()

    pdf.set_font("Times", "B", 11)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(col_w, 8, "CORE COMPETENCIES", new_x="RIGHT", new_y="TOP")
    pdf.cell(8, 8, "")
    pdf.cell(col_w, 8, "COMPANY DATA", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Times", "", 10)
    y_left = pdf.get_y()
    pdf.multi_cell(col_w, 5, (cap_body or "Facilities support services.")[:800])
    y_after_left = pdf.get_y()
    pdf.set_xy(pdf.l_margin + col_w + 8, y_left)
    lines = [
        f"CAGE: {owner.get('cage_code') or '-'}",
        f"UEI: {owner.get('uei') or '-'}",
        f"EIN: {owner.get('ein') or '-'}",
        f"NAICS: {owner.get('primary_naics_code', '561720')}",
        f"Phone: {owner.get('business_phone', '')}",
        f"Email: {owner.get('business_email', '')}",
    ]
    pdf.multi_cell(col_w, 5, "\n".join(lines))
    y_after_right = pdf.get_y()
    pdf.set_y(max(y_after_left, y_after_right) + 6)

    pdf.set_font("Times", "B", 11)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(col_w, 8, "PAST PERFORMANCE", new_x="RIGHT", new_y="TOP")
    pdf.cell(8, 8, "")
    pdf.cell(col_w, 8, "CONTACT", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Times", "", 10)
    y_left = pdf.get_y()
    pdf.multi_cell(col_w, 5, (past_body or owner.get("commercial_experience") or "")[:600])
    pdf.set_xy(pdf.l_margin + col_w + 8, y_left)
    pdf.multi_cell(
        col_w,
        5,
        f"{owner.get('owner_name', 'Mark Graham II')}, {owner.get('owner_title', 'Owner')}",
    )

    return _fpdf_bytes(pdf)


def build_capability_pdf(proposal: Any, sections: dict[str, str], meta: dict[str, Any]) -> tuple[bytes, str]:
    html = _capability_html(proposal, sections, meta)
    if _weasyprint_available():
        try:
            from weasyprint import HTML

            return HTML(string=html).write_pdf(optimize_size=("fonts", "images")), "weasyprint"
        except Exception:
            pass
    return build_capability_pdf_fpdf(proposal, sections, meta), "fpdf2"
